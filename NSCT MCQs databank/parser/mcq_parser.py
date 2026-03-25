#!/usr/bin/env python3
"""
NSCT MCQ Parser
Converts all MCQ files to unified JSON matching the DB contract:
  option_a = correct answer, option_b-e = wrong answers (nullable)

Usage:
  python3 mcq_parser.py                  # parse all -> output/all_questions.json
  python3 mcq_parser.py --file <path>    # parse single file
  python3 mcq_parser.py --stats          # stats only, no file output
"""

import re, json, argparse, sys, subprocess
from pathlib import Path

try:
    import docx as _docx
except ImportError:
    _docx = None
try:
    import pdfplumber as _pdf
except ImportError:
    _pdf = None
try:
    import openpyxl as _xl
except ImportError:
    _xl = None

DATABANK_ROOT = Path(__file__).parent.parent

# ---------------------------------------------------------------------------
# Noise patterns — lines to discard
# ---------------------------------------------------------------------------
NOISE_RE = re.compile(
    r'^(department[\s:]*|program[\s:]*|subject\s*(name|title|code)[\s:]*'
    r'|course[\s:]*|java\s*programming\s*$|python\s*programming.*$'
    r'|mcqs?\s+for\s+|answer\s*key[s]?.*|section\s+\d+.*|topic\s+\d+.*'
    r'|_{3,}|-{3,}|={3,})',
    re.IGNORECASE
)

# ---------------------------------------------------------------------------
# Option prefix  A. A) a. a)  i) ii) iii) iv) v)  bullet "o "
# ---------------------------------------------------------------------------
OPTION_RE = re.compile(
    r'^(?:([A-Ea-e])[.)]\s*|([ivxIVX]+)[.)]\s*|o\s+)'
)
ROMAN_MAP = {'i': 0, 'ii': 1, 'iii': 2, 'iv': 3, 'v': 4}

# ---------------------------------------------------------------------------
# Answer line  ANSWER: B  ANSWER B  Answer: ii) text  Correct Answer: ii) text
# ---------------------------------------------------------------------------
ANSWER_RE = re.compile(
    r'^(?:correct\s+)?answer\s*:?\s*(?:([A-Ea-e])\b|([ivxIVX]+)[.)]\s*(.*))',
    re.IGNORECASE
)

QNUM_RE = re.compile(r'^\d+[.)]\s+')


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def parse_option_prefix(line):
    """Returns (index 0-4, text) or None. index=None for bullet 'o'."""
    m = OPTION_RE.match(line)
    if not m:
        return None
    letter, roman = m.group(1), m.group(2)
    text = line[m.end():]
    if letter:
        return ord(letter.upper()) - ord('A'), text.strip()
    if roman:
        idx = ROMAN_MAP.get(roman.lower())
        return (idx, text.strip()) if idx is not None else None
    return None, text.strip()   # bullet


def parse_answer_line(line):
    """Returns (0-based index, answer_text_or_None) or (None, None)."""
    m = ANSWER_RE.match(line.strip())
    if not m:
        return None, None
    letter, roman, roman_text = m.group(1), m.group(2), m.group(3)
    if letter:
        return ord(letter.upper()) - ord('A'), None
    if roman:
        return ROMAN_MAP.get(roman.lower()), (roman_text or '').strip()
    return None, None


def strip_qnum(line):
    return QNUM_RE.sub('', line).strip()


def assemble(question_text, options, answer_idx, source_file, subject):
    """Rotate options so correct answer becomes option_a."""
    question_text = question_text.strip()
    options = [o.strip() for o in options if o and o.strip()]
    if not question_text or len(options) < 2:
        return None
    if answer_idx is None or answer_idx >= len(options):
        return None
    correct = options[answer_idx]
    wrongs  = [o for i, o in enumerate(options) if i != answer_idx]
    return {
        'subject':       subject,
        'question_text': question_text,
        'option_a':      correct,
        'option_b':      wrongs[0] if len(wrongs) > 0 else None,
        'option_c':      wrongs[1] if len(wrongs) > 1 else None,
        'option_d':      wrongs[2] if len(wrongs) > 2 else None,
        'option_e':      wrongs[3] if len(wrongs) > 3 else None,
        'source_file':   source_file,
    }


# ---------------------------------------------------------------------------
# FORMAT A — All-in-one paragraphs + trailing answer key
# Files: AI ML DA.docx, Problem Solving.docx
# Each docx paragraph = full question + options joined by \n
# Answer key section at end: one letter per line after "ANSWER KEY" header
# ---------------------------------------------------------------------------
def parse_allinone_with_answerkey(lines, source_file, subject):
    ak_start = None
    for i, l in enumerate(lines):
        if re.match(r'^answer\s*keys?\s*$', l.strip(), re.IGNORECASE):
            ak_start = i
            break
    if ak_start is None:
        return []

    answer_letters = []
    for l in lines[ak_start + 1:]:
        t = l.strip()
        if re.match(r'^[A-Ea-e]$', t):
            answer_letters.append(ord(t.upper()) - ord('A'))

    results = []
    q_idx   = 0

    for line in lines[:ak_start]:
        line = line.strip()
        if not line or NOISE_RE.match(line):
            continue
        # Question block: embedded \n with option markers
        if '\n' in line and re.search(r'\n[A-Ea-e][.)]\s', line):
            parts      = line.split('\n')
            q_text     = strip_qnum(parts[0].strip())
            options    = []
            bullet_seq = 0
            for part in parts[1:]:
                part = part.strip()
                if not part:
                    continue
                parsed = parse_option_prefix(part)
                if parsed:
                    idx, text = parsed
                    if idx is None:
                        idx = bullet_seq
                        bullet_seq += 1
                    while len(options) <= idx:
                        options.append('')
                    options[idx] = text
            if q_idx < len(answer_letters) and len(options) >= 2:
                rec = assemble(q_text, options, answer_letters[q_idx], source_file, subject)
                if rec:
                    results.append(rec)
            q_idx += 1

    return results


# ---------------------------------------------------------------------------
# FORMAT B — Generic line-by-line with inline ANSWER
# Files: OS, Java, Python, DSA, Cyber Security, Database, SE, MCQ Data Structures
# ---------------------------------------------------------------------------
def parse_lines(lines, source_file, subject):
    results    = []
    q_text     = None
    options    = []
    bullet_seq = 0
    answer_idx = None

    def flush():
        nonlocal q_text, options, bullet_seq, answer_idx
        if q_text and options and answer_idx is not None:
            rec = assemble(q_text, options, answer_idx, source_file, subject)
            if rec:
                results.append(rec)
        q_text = None; options = []; bullet_seq = 0; answer_idx = None

    # Expand any paragraph that packs multiple options with \n
    expanded = []
    for raw in lines:
        if '\n' in raw:
            expanded.extend(raw.split('\n'))
        else:
            expanded.append(raw)

    for raw in expanded:
        line = raw.strip()
        if not line or NOISE_RE.match(line):
            continue

        # Answer line
        ans_idx, ans_text = parse_answer_line(line)
        if ans_idx is not None:
            if ans_text and options:   # roman-numeral: match by text prefix
                for i, opt in enumerate(options):
                    if opt.lower().startswith(ans_text.lower()[:20]):
                        ans_idx = i; break
            answer_idx = ans_idx
            flush()
            continue

        # Explanation line (Database PDF)
        if re.match(r'^explanation\s*:', line, re.IGNORECASE):
            continue

        # Option line
        parsed = parse_option_prefix(line)
        if parsed is not None:
            idx, text = parsed
            if idx is None:
                idx = bullet_seq
                bullet_seq += 1
            while len(options) <= idx:
                options.append('')
            options[idx] = text
            continue

        # New question — discard pending unanswered (e.g. Cloud Computing)
        if q_text and options and answer_idx is None:
            q_text = None; options = []; bullet_seq = 0

        q_text = strip_qnum(line); options = []; bullet_seq = 0; answer_idx = None

    if answer_idx is not None:
        flush()
    return results


# ---------------------------------------------------------------------------
# FORMAT C — Web Dev: questions file + separate answer key file
# Questions: alternating paragraphs (question text, then options block with \n)
# Answer key: one letter per line
# ---------------------------------------------------------------------------
def parse_webdev(questions_path, answerkey_path, subject):
    if _docx is None:
        raise RuntimeError("python-docx not installed")

    ak_doc = _docx.Document(str(answerkey_path))
    answer_letters = []
    for p in ak_doc.paragraphs:
        t = p.text.strip()
        if re.match(r'^[A-Da-d]$', t):
            answer_letters.append(ord(t.upper()) - ord('A'))

    q_doc   = _docx.Document(str(questions_path))
    paras   = [p.text.strip() for p in q_doc.paragraphs if p.text.strip()]
    results = []
    q_idx   = 0
    i       = 0

    while i < len(paras):
        para = paras[i]
        if NOISE_RE.match(para) or re.match(r'^mcqs?\s+for\s+', para, re.IGNORECASE):
            i += 1; continue
        if '\n' in para and re.match(r'^[A-Da-d][.)]\s', para):
            i += 1; continue
        if i + 1 < len(paras) and '\n' in paras[i + 1]:
            q_text  = strip_qnum(para)
            options = []
            for opt in paras[i + 1].split('\n'):
                parsed = parse_option_prefix(opt.strip())
                if parsed:
                    _, text = parsed
                    options.append(text)
            if q_idx < len(answer_letters) and len(options) >= 2:
                rec = assemble(q_text, options, answer_letters[q_idx],
                               questions_path.name, subject)
                if rec:
                    results.append(rec)
            q_idx += 1; i += 2; continue
        i += 1

    return results


# ---------------------------------------------------------------------------
# FORMAT D — Excel (Computer Networks)
# Columns: Question | Correct | Wrong1 | Wrong2 | Wrong3
# Already in DB contract — no rotation needed
# ---------------------------------------------------------------------------
def parse_xlsx(path, subject):
    if _xl is None:
        raise RuntimeError("openpyxl not installed")
    wb = _xl.load_workbook(str(path))
    ws = wb.active
    results = []
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        if i == 0:
            continue
        if not row[0]:
            continue
        q_text, correct, w1, w2, w3 = (str(c).strip() if c else None for c in row[:5])
        if not q_text or not correct:
            continue
        results.append({
            'subject': subject, 'question_text': q_text,
            'option_a': correct, 'option_b': w1, 'option_c': w2,
            'option_d': w3, 'option_e': None, 'source_file': path.name,
        })
    return results


# ---------------------------------------------------------------------------
# Text extractors
# ---------------------------------------------------------------------------
def extract_docx(path):
    if _docx is None:
        raise RuntimeError("python-docx not installed")
    doc = _docx.Document(str(path))
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def extract_doc(path):
    result = subprocess.run(
        ['textutil', '-convert', 'txt', '-stdout', str(path)],
        capture_output=True, text=True
    )
    if result.returncode != 0:
        raise RuntimeError(f"textutil failed: {result.stderr}")
    return result.stdout.split('\n')


def extract_pdf(path):
    if _pdf is None:
        raise RuntimeError("pdfplumber not installed")
    lines = []
    with _pdf.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                lines.extend(text.split('\n'))
    return lines


# ---------------------------------------------------------------------------
# File registry
# ---------------------------------------------------------------------------
def build_registry(root):
    R = []

    def reg(rel, extractor, subject, parser='lines', extra=None):
        R.append({'path': root / rel, 'extractor': extractor,
                  'subject': subject, 'parser': parser, 'extra': extra})

    # Format A
    reg('AI ML DA.docx',
        'docx', 'AI / ML / Data Analytics', parser='allinone')
    reg('Problem Solving and Analytical Skills (1).docx',
        'docx', 'Problem Solving & Analytical Skills', parser='allinone')

    # Format B
    reg('Operating System/Part II - Operating System.docx',  'docx', 'Operating System')
    reg('Programming/Java Programming - Part I.docx',        'docx', 'Java Programming')
    reg('Programming/Java Programming - Part II.docx',       'docx', 'Java Programming')
    reg('Programming/Java Programming - Part III.docx',      'docx', 'Java Programming')
    reg('Programming/Python Programming- Part I.docx',       'docx', 'Python Programming')
    reg('Programming/Python Programming- Part II.docx',      'docx', 'Python Programming')
    reg('Operating System/Part I - Operating System.doc',    'doc',  'Operating System')
    reg('Data Strucutres & Algorithms/PART 1(BSCS-CSIT-01306-Data Structure & Algorithms-03-Computer Science & IT).doc',
        'doc', 'Data Structures & Algorithms')
    reg('Data Strucutres & Algorithms/PART 2(BSCS-CSIT-01306-Data Structure & Algorithms-03-Computer Science & IT).doc',
        'doc', 'Data Structures & Algorithms')
    reg('Cyber Security 300 MCQs Bank.pdf',                  'pdf',  'Cyber Security')
    reg('Database MCQs.pdf',                                 'pdf',  'Database Systems')
    reg('MCQ bank Data Structures.pdf',                      'pdf',  'Data Structures & Algorithms')
    reg('Software Engineering 200 MCQs with Answer.pdf',     'pdf',  'Software Engineering')

    # Format C
    R.append({
        'path':      root / 'Web Development/MCQS for Web Development.docx',
        'extractor': 'webdev', 'subject': 'Web Development',
        'parser':    'webdev',
        'extra':     root / 'Web Development/Answer Key.docx',
    })

    # Format D
    R.append({
        'path':      root / 'Computer Networks MCQs.xlsx',
        'extractor': 'xlsx', 'subject': 'Computer Networks',
        'parser':    'xlsx', 'extra': None,
    })

    # Skipped:
    #   Software Engineering 200 MCQs with Answer (1).pdf  — exact duplicate
    #   1773309682538.pdf                                   — image-based lecture notes
    #   MCQ bank Cloud Computing.pdf                        — no answer keys in file

    return R


# ---------------------------------------------------------------------------
# Orchestrator
# ---------------------------------------------------------------------------
def parse_file(entry):
    path    = entry['path']
    subject = entry['subject']
    parser  = entry['parser']
    ext     = entry['extractor']

    if not path.exists():
        print(f"  [SKIP] not found: {path.name}", file=sys.stderr)
        return []

    try:
        if parser == 'xlsx':
            return parse_xlsx(path, subject)
        if parser == 'webdev':
            return parse_webdev(path, entry['extra'], subject)

        if ext == 'docx':
            lines = extract_docx(path)
        elif ext == 'doc':
            lines = extract_doc(path)
        elif ext == 'pdf':
            lines = extract_pdf(path)
        else:
            print(f"  [SKIP] unknown extractor: {ext}", file=sys.stderr)
            return []

        if parser == 'allinone':
            return parse_allinone_with_answerkey(lines, path.name, subject)

        return parse_lines(lines, path.name, subject)

    except Exception as e:
        print(f"  [ERROR] {path.name}: {e}", file=sys.stderr)
        return []


def run(target_path=None, stats_only=False):
    registry      = build_registry(DATABANK_ROOT)
    all_questions = []

    if target_path:
        tp       = Path(target_path).resolve()
        registry = [e for e in registry if e['path'].resolve() == tp]
        if not registry:
            print(f"File not in registry: {target_path}")
            return []

    for entry in registry:
        name = entry['path'].name
        print(f"Parsing: {name} ...", end=' ', flush=True)
        questions = parse_file(entry)
        print(f"{len(questions)} questions")
        all_questions.extend(questions)

    print(f"\nTotal: {len(all_questions)} questions parsed")

    if stats_only:
        return all_questions

    out_dir = Path(__file__).parent / 'output'
    out_dir.mkdir(exist_ok=True)

    out_path = out_dir / 'all_questions.json'
    with open(out_path, 'w', encoding='utf-8') as f:
        json.dump(all_questions, f, ensure_ascii=False, indent=2)
    print(f"\nWritten -> {out_path}")

    by_subject = {}
    for q in all_questions:
        by_subject.setdefault(q['subject'], []).append(q)

    for subj, qs in sorted(by_subject.items()):
        safe = re.sub(r'[^\w\s-]', '', subj).strip().replace(' ', '_')
        sp   = out_dir / f"{safe}.json"
        with open(sp, 'w', encoding='utf-8') as f:
            json.dump(qs, f, ensure_ascii=False, indent=2)
        print(f"  {subj}: {len(qs)} -> {sp.name}")

    return all_questions


if __name__ == '__main__':
    ap = argparse.ArgumentParser(description='NSCT MCQ Parser')
    ap.add_argument('--file',  help='Parse a single file by path')
    ap.add_argument('--stats', action='store_true', help='Stats only, no file output')
    args = ap.parse_args()
    run(target_path=args.file, stats_only=args.stats)
